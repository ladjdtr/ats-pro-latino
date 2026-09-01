import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
import requests
import trafilatura
from io import BytesIO
from docx import Document
from fpdf import FPDF
import os
import time
from dotenv import load_dotenv

load_dotenv()

# Configurar Gemini (si no tiene API Key, le pedimos una)
API_KEY = os.getenv("GEMINI_API_KEY", "")
if API_KEY:
    genai.configure(api_key=API_KEY)

# --- FUNCIONES ---
def extraer_texto_cv(uploaded_file, texto_plano=""):
    if texto_plano and texto_plano.strip() != "":
        return texto_plano
    if uploaded_file is None:
        return ""
    try:
        file_bytes = uploaded_file.read()
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(BytesIO(file_bytes))
            return "\n".join([page.extract_text() or "" for page in reader.pages])
        elif "wordprocessingml" in uploaded_file.type:
            doc = docx.Document(BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return "Formato no soportado. Usa PDF o Word."
    except Exception as e:
        return f"Error: {e}"

def extraer_texto_oferta(url):
    if not url or url.strip() == "":
        return None
    try:
        response = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=10)
        if response.status_code == 200:
            texto = trafilatura.extract(response.text, include_comments=False, include_tables=True)
            return texto if texto else None
        return None
    except:
        return None

def generar_adaptacion(cv_texto, job_texto, api_key):
    if not api_key:
        return "Error: Necesitas una API Key de Google Gemini."
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-1.5-pro')
    
    prompt = f"""Eres un reclutador canadiense senior especializado en ATS.
    Reglas:
    1. Devuelve SOLO el contenido, sin introducciones.
    2. SEPARA el CV y la Cover Letter con estas etiquetas exactas:
    ---RESUME---
    [Aquí el CV adaptado]
    ---COVER---
    [Aquí la Cover Letter en formato email]

    CV del usuario:
    {cv_texto}

    Oferta de trabajo:
    {job_texto}
    """
    response = model.generate_content(prompt)
    return response.text

def generar_word(texto):
    doc = Document()
    for line in texto.split('\n'):
        doc.add_paragraph(line.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generar_pdf(texto):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in texto.split('\n'):
        line = line.encode('latin-1', 'ignore').decode('latin-1')
        pdf.cell(200, 10, txt=line[:80], ln=True, align='L')
    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFAZ STREAMLIT ---
st.set_page_config(page_title="ATS Pro Latino", layout="centered")
st.title("🚀 ATS Pro Latino")
st.caption("Adapta tu CV al mercado Canadiense/USA. ¡Pasa los filtros ATS!")

# Sidebar
with st.sidebar:
    st.header("⚙️ Configuración")
    user_api_key = st.text_input("🔑 API Key de Gemini (opcional)", type="password")
    st.divider()
    st.markdown("**Plan Free:** 1 CV al día (sesión local)")
    st.markdown("**Plan Pro:** $9.99 CAD/mes (Próximamente)")

# Estado de uso (simple local)
if 'contador' not in st.session_state:
    st.session_state.contador = 0
if 'fecha' not in st.session_state:
    st.session_state.fecha = time.strftime("%Y-%m-%d")

if st.session_state.fecha != time.strftime("%Y-%m-%d"):
    st.session_state.contador = 0
    st.session_state.fecha = time.strftime("%Y-%m-%d")

# Entrada de datos
col1, col2 = st.columns(2)
with col1:
    st.subheader("📄 Tu CV")
    uploaded = st.file_uploader("Sube PDF/Word", type=["pdf", "docx"])
    cv_text = st.text_area("O pégalo aquí", height=100)

with col2:
    st.subheader("💼 Oferta")
    job_url = st.text_input("🔗 URL (Indeed, Zip, JobBank)")
    job_text = st.text_area("O pega la descripción", height=150)

if st.button("✨ Adaptar mi CV ya! (1 uso/día)"):
    if st.session_state.contador >= 1:
        st.error("🔥 Límite diario alcanzado. ¡Hazte Pro pronto!")
    else:
        cv_final = extraer_texto_cv(uploaded, cv_text)
        job_final = job_text
        if job_url:
            scraped = extraer_texto_oferta(job_url)
            if scraped:
                job_final = scraped
            else:
                st.warning("No pude leer el link automáticamente.")

        if len(cv_final) < 20 or len(job_final) < 20:
            st.error("Falta el CV o la oferta.")
        else:
            api_key = user_api_key if user_api_key else API_KEY
            if not api_key:
                st.error("Falta la API Key de Gemini. Ponla en la barra lateral.")
            else:
                with st.spinner("Reclutador canadiense analizando tu perfil..."):
                    resultado = generar_adaptacion(cv_final, job_final, api_key)
                
                st.session_state.contador += 1
                st.success("¡Listo!")
                
                # Separar usando las etiquetas ---RESUME--- y ---COVER---
                if "---COVER---" in resultado:
                    partes = resultado.split("---COVER---")
                    cv_out = partes[0].replace("---RESUME---", "").strip()
                    cover_out = partes[1].strip()
                else:
                    # Fallback: si no encuentra las etiquetas, muestra todo como CV
                    cv_out = resultado
                    cover_out = "No se pudo generar Cover Letter. Intenta de nuevo."
                
                tab1, tab2, tab3 = st.tabs(["📝 CV Adaptado", "✉️ Cover Letter", "⬇️ Descargar"])
                with tab1:
                    st.text_area("CV Adaptado", cv_out, height=350)
                with tab2:
                    st.text_area("Cover Letter", cover_out, height=200)
                with tab3:
                    # Descarga en WORD
                    st.download_button(
                        "📥 Descargar CV + Cover en Word",
                        data=generar_word(resultado),
                        file_name="CV_Cover_Adaptado.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    # Descarga en PDF
                    st.download_button(
                        "📥 Descargar CV + Cover en PDF",
                        data=generar_pdf(resultado),
                        file_name="CV_Cover_Adaptado.pdf",
                        mime="application/pdf"
                    )
                    # Solo Cover en texto plano (para email)
                    st.download_button(
                        "✉️ Solo Cover (Texto para email)",
                        data=cover_out,
                        file_name="cover_letter.txt",
                        mime="text/plain"
                    )
